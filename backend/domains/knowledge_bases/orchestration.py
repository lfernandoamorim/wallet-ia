"""Camada de Orquestração para o domínio de Base de Conhecimento e RAG."""

import hashlib
import io
import math
import os
import re
import uuid
from typing import Any
import docx
import openpyxl
from fastapi import HTTPException, UploadFile, status
from sqlalchemy import or_, select
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy.orm import selectinload

from backend.core.config import settings
from backend.domains.knowledge_bases.execution import KBChunk, KBDocument, KnowledgeBase
from backend.domains.users.execution import User


def extract_text_from_file(file_name: str, content_bytes: bytes) -> str:
    """Extrai texto legível de arquivos .md, .txt, .docx e .xlsx."""
    ext = file_name.split(".")[-1].lower()

    if ext in ["md", "txt", "markdown", "csv"]:
        return content_bytes.decode("utf-8", errors="ignore")

    elif ext in ["docx", "doc"]:
        try:
            doc = docx.Document(io.BytesIO(content_bytes))
            full_text = []
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text.append(para.text.strip())
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_text:
                        full_text.append(" | ".join(row_text))
            return "\n\n".join(full_text)
        except Exception as e:
            return f"[Erro ao extrair docx: {str(e)}]"

    elif ext in ["xlsx", "xls"]:
        try:
            wb = openpyxl.load_workbook(io.BytesIO(content_bytes), data_only=True)
            lines = []
            for sheet in wb.sheetnames:
                ws = wb[sheet]
                lines.append(f"--- Planilha: {sheet} ---")
                for row in ws.iter_rows(values_only=True):
                    row_values = [str(val).strip() for val in row if val is not None and str(val).strip()]
                    if row_values:
                        lines.append(" | ".join(row_values))
            return "\n".join(lines)
        except Exception as e:
            return f"[Erro ao extrair xlsx: {str(e)}]"

    return content_bytes.decode("utf-8", errors="ignore")


def split_text_into_chunks(text: str, chunk_size: int = 600, overlap: int = 100) -> list[str]:
    """Divide um texto em pedaços (chunks) com sobreposição para contexto."""
    cleaned = re.sub(r"\s+", " ", text).strip()
    if not cleaned:
        return []

    chunks = []
    start = 0
    while start < len(cleaned):
        end = start + chunk_size
        chunk = cleaned[start:end]
        chunks.append(chunk)
        start += chunk_size - overlap
        if end >= len(cleaned):
            break
    return chunks


def generate_embedding_mock(text: str, dimension: int = 1536) -> list[float]:
    """
    Gera um vetor semântico normalizado determinístico baseado em hash para ambientes
    locais e testes (em produção, pode ser alimentado via OpenAI text-embedding-3-small).
    """
    h = hashlib.sha256(text.encode("utf-8")).digest()
    raw = [float(b) / 255.0 for b in h]
    # Expande repetindo até atingir a dimensão requerida
    vector = (raw * (dimension // len(raw) + 1))[:dimension]
    # Normalização L2
    norm = math.sqrt(sum(x * x for x in vector)) or 1.0
    return [x / norm for x in vector]


def cosine_similarity(vec_a: list[float], vec_b: list[float]) -> float:
    """Calcula a similaridade por cosseno entre dois vetores."""
    dot = sum(a * b for a, b in zip(vec_a, vec_b))
    norm_a = math.sqrt(sum(a * a for a in vec_a))
    norm_b = math.sqrt(sum(b * b for b in vec_b))
    if norm_a == 0.0 or norm_b == 0.0:
        return 0.0
    return dot / (norm_a * norm_b)


async def create_knowledge_base(
    session: AsyncSession,
    owner_id: str,
    name: str,
    description: str | None = None,
    visibility: str = "private",
) -> KnowledgeBase:
    """Cria uma nova base de conhecimento."""
    slug = str(uuid.uuid4())[:8] if visibility == "public" else None
    kb = KnowledgeBase(
        owner_id=owner_id,
        name=name,
        description=description,
        visibility=visibility,
        public_slug=slug,
    )
    session.add(kb)
    await session.commit()
    await session.refresh(kb)
    return kb


async def list_knowledge_bases(
    session: AsyncSession,
    user_id: str,
    can_view_all: bool = False,
) -> list[KnowledgeBase]:
    """Lista as bases de conhecimento que o usuário tem permissão para visualizar."""
    if can_view_all:
        query = select(KnowledgeBase).order_by(KnowledgeBase.created_at.desc())
    else:
        query = select(KnowledgeBase).where(
            or_(KnowledgeBase.owner_id == user_id, KnowledgeBase.visibility == "public")
        ).order_by(KnowledgeBase.created_at.desc())
    result = await session.execute(query)
    return list(result.scalars().all())


async def get_knowledge_base_by_id(
    session: AsyncSession,
    kb_id: str,
    user_id: str | None = None,
    can_view_all: bool = False,
) -> KnowledgeBase:
    """Busca uma base de conhecimento por ID validando visibilidade."""
    query = select(KnowledgeBase).where(KnowledgeBase.id == kb_id).options(
        selectinload(KnowledgeBase.documents)
    )
    result = await session.execute(query)
    kb = result.scalars().first()
    if not kb:
        raise HTTPException(status_code=404, detail="Base de conhecimento não encontrada.")

    if not can_view_all and kb.visibility != "public" and user_id and str(kb.owner_id) != str(user_id):
        raise HTTPException(status_code=403, detail="Acesso não autorizado a esta base de conhecimento.")

    return kb


async def ingest_document(
    session: AsyncSession,
    knowledge_base_id: str,
    file_name: str,
    content_bytes: bytes,
    uploaded_by_id: str,
) -> KBDocument:
    """Processa um arquivo, extrai texto, divide em chunks e gera embeddings para RAG."""
    os.makedirs(settings.storage_path, exist_ok=True)
    saved_file_name = f"{uuid.uuid4()}_{file_name}"
    storage_path = os.path.join(settings.storage_path, saved_file_name)
    with open(storage_path, "wb") as f:
        f.write(content_bytes)

    file_ext = file_name.split(".")[-1].lower() if "." in file_name else "txt"
    doc = KBDocument(
        knowledge_base_id=knowledge_base_id,
        file_name=file_name,
        file_type=file_ext,
        storage_path=storage_path,
        status="processing",
        uploaded_by=uploaded_by_id,
    )
    session.add(doc)
    await session.flush()

    try:
        extracted_text = extract_text_from_file(file_name, content_bytes)
        chunks = split_text_into_chunks(extracted_text)

        for idx, chunk_text in enumerate(chunks):
            embedding = generate_embedding_mock(chunk_text)
            kb_chunk = KBChunk(
                document_id=doc.id,
                knowledge_base_id=knowledge_base_id,
                content=chunk_text,
                embedding=embedding,
                chunk_index=idx,
            )
            session.add(kb_chunk)

        doc.status = "ready"
    except Exception as e:
        doc.status = "error"
        doc.error_message = str(e)

    await session.commit()
    await session.refresh(doc)
    return doc


async def retrieve_relevant_chunks(
    session: AsyncSession,
    knowledge_base_ids: list[str],
    query_text: str,
    top_k: int = 5,
) -> list[str]:
    """Busca os trechos mais relevantes do RAG dentro das bases de conhecimento especificadas."""
    if not knowledge_base_ids:
        return []

    query = select(KBChunk).where(KBChunk.knowledge_base_id.in_(knowledge_base_ids))
    result = await session.execute(query)
    chunks = list(result.scalars().all())

    if not chunks:
        return []

    query_vec = generate_embedding_mock(query_text)
    scored_chunks = []
    for c in chunks:
        if c.embedding is not None:
            # Converte de vetor do pgvector ou lista
            vec = list(c.embedding) if hasattr(c.embedding, "__iter__") else []
            sim = cosine_similarity(query_vec, vec) if vec else 0.0
            scored_chunks.append((sim, c.content))

    scored_chunks.sort(key=lambda x: x[0], reverse=True)
    return [content for _, content in scored_chunks[:top_k]]
